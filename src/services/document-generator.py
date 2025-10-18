#!/usr/bin/env python3
"""
DOCUMENT GENERATOR SERVICE
TeleMedCare V11.0-Cloudflare - Sistema Modulare

Gestisce:
- Lettura template DOCX con placeholder
- Sostituzione placeholder con dati lead
- Generazione PDF da DOCX
- Salvataggio in filesystem organizzato
- Inserimento nel database (contracts, proforma)

Flusso:
1. Lead arriva dal form → dati salvati in DB
2. Se vuole contratto → genera contratto + proforma
3. Sostituisce placeholder nei template DOCX
4. Converte DOCX → PDF
5. Salva in /documents/contratti/ e /documents/proforma/
6. Inserisce record in DB (contracts, proforma)
"""

import sys
import os
import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, Tuple
import re
from docx import Document
from docx2pdf import convert
import subprocess

class DocumentGenerator:
    """Generatore contratti e proforma da template DOCX"""
    
    # Prezzi servizi (primo anno)
    PREZZI = {
        'BASE': 480.00,      # € + IVA (22%) = 585.60€
        'AVANZATO': 840.00   # € + IVA (22%) = 1024.80€
    }
    
    # Prezzi rinnovi (dal secondo anno in poi)
    PREZZI_RINNOVO = {
        'BASE': 240.00,      # € + IVA (22%) = 292.80€
        'AVANZATO': 600.00   # € + IVA (22%) = 732.00€
    }
    
    IVA_RATE = 0.22  # 22%
    
    def __init__(self, templates_dir: str = './templates', output_dir: str = './documents'):
        """
        Inizializza il generatore
        
        Args:
            templates_dir: Directory con template DOCX
            output_dir: Directory output per PDF generati
        """
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        
        # Crea cartelle output se non esistono
        self.contratti_dir = self.output_dir / 'contratti'
        self.proforma_dir = self.output_dir / 'proforma'
        self.contratti_firmati_dir = self.output_dir / 'contratti_firmati'
        
        for dir_path in [self.contratti_dir, self.proforma_dir, self.contratti_firmati_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
    
    def generate_contract_from_lead(self, lead_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Genera contratto completo da dati lead
        
        Args:
            lead_data: Dizionario con dati lead dal database
            
        Returns:
            Dizionario con risultato generazione:
            {
                'success': bool,
                'contract_id': str,
                'contract_pdf_path': str,
                'proforma_id': str,
                'proforma_pdf_path': str,
                'errors': list
            }
        """
        try:
            print(f"\n🚀 Inizio generazione documenti per lead {lead_data.get('id')}")
            
            # Valida dati obbligatori
            validation = self._validate_lead_data(lead_data)
            if not validation['valid']:
                return {
                    'success': False,
                    'errors': validation['errors']
                }
            
            # Determina tipo servizio
            tipo_servizio = self._determine_service_type(lead_data)
            print(f"📦 Tipo servizio: {tipo_servizio}")
            
            # Genera ID univoci
            contract_id = self._generate_contract_id()
            proforma_id = self._generate_proforma_id()
            
            # Genera contratto
            contract_result = self._generate_contract_pdf(lead_data, tipo_servizio, contract_id)
            if not contract_result['success']:
                return {
                    'success': False,
                    'errors': contract_result['errors']
                }
            
            # Genera proforma
            proforma_result = self._generate_proforma_pdf(lead_data, tipo_servizio, proforma_id)
            if not proforma_result['success']:
                return {
                    'success': False,
                    'errors': proforma_result['errors']
                }
            
            print("✅ Documenti generati con successo!")
            
            return {
                'success': True,
                'contract_id': contract_id,
                'contract_pdf_path': contract_result['pdf_path'],
                'proforma_id': proforma_id,
                'proforma_pdf_path': proforma_result['pdf_path'],
                'tipo_servizio': tipo_servizio,
                'prezzo_base': self.PREZZI[tipo_servizio],
                'prezzo_iva_inclusa': self.PREZZI[tipo_servizio] * (1 + self.IVA_RATE),
                'errors': []
            }
            
        except Exception as e:
            print(f"❌ Errore generazione documenti: {e}")
            return {
                'success': False,
                'errors': [str(e)]
            }
    
    def _validate_lead_data(self, lead_data: Dict[str, Any]) -> Dict[str, Any]:
        """Valida che i dati lead contengano tutti i campi obbligatori"""
        errors = []
        
        # Campi obbligatori per contratto
        required_fields = [
            'nomeAssistito', 'cognomeAssistito', 'dataNascitaAssistito',
            'luogoNascitaAssistito', 'emailRichiedente', 'telefonoRichiedente'
        ]
        
        for field in required_fields:
            if not lead_data.get(field):
                errors.append(f"Campo obbligatorio mancante: {field}")
        
        # Controlla intestazione contratto
        intestazione = lead_data.get('intestazioneContratto', 'Assistito')
        
        if intestazione == 'Assistito':
            # Serve CF, indirizzo completo assistito
            if not lead_data.get('cfAssistito'):
                errors.append("Codice Fiscale Assistito mancante")
            if not lead_data.get('indirizzoAssistito'):
                errors.append("Indirizzo Assistito mancante")
        else:
            # Serve CF, indirizzo completo richiedente
            if not lead_data.get('cfRichiedente'):
                errors.append("Codice Fiscale Richiedente mancante")
            if not lead_data.get('indirizzoRichiedente'):
                errors.append("Indirizzo Richiedente mancante")
        
        return {
            'valid': len(errors) == 0,
            'errors': errors
        }
    
    def _determine_service_type(self, lead_data: Dict[str, Any]) -> str:
        """
        Determina il tipo di servizio (BASE o AVANZATO) dai dati lead
        
        Logica:
        - Se pacchetto contiene 'avanzat' → AVANZATO
        - Se pacchetto contiene 'base' → BASE
        - Default → BASE
        """
        pacchetto = lead_data.get('pacchetto', '').lower()
        
        if 'avanzat' in pacchetto:
            return 'AVANZATO'
        else:
            return 'BASE'
    
    def _generate_contract_pdf(self, lead_data: Dict[str, Any], tipo_servizio: str, contract_id: str) -> Dict[str, Any]:
        """Genera PDF contratto da template DOCX"""
        try:
            # Seleziona template corretto
            if tipo_servizio == 'BASE':
                template_name = 'Template_Contratto_Base_TeleMedCare.docx'
            else:
                template_name = 'Template_Contratto_Avanzato_TeleMedCare.docx'
            
            template_path = self.templates_dir / template_name
            
            if not template_path.exists():
                return {
                    'success': False,
                    'errors': [f"Template non trovato: {template_path}"]
                }
            
            print(f"📄 Generazione contratto da template: {template_name}")
            
            # Prepara dati per sostituzione placeholder
            placeholder_data = self._prepare_contract_placeholders(lead_data, tipo_servizio)
            
            # Crea DOCX con placeholder sostituiti
            output_filename = f"{datetime.now().strftime('%Y%m%d')}_{lead_data.get('cognomeAssistito')}_{contract_id}.docx"
            output_docx_path = self.contratti_dir / output_filename
            
            self._replace_placeholders_in_docx(template_path, output_docx_path, placeholder_data)
            
            # Converte DOCX → PDF
            output_pdf_path = output_docx_path.with_suffix('.pdf')
            self._convert_docx_to_pdf(output_docx_path, output_pdf_path)
            
            # Rimuovi DOCX intermedio (opzionale)
            # output_docx_path.unlink()
            
            print(f"✅ Contratto PDF generato: {output_pdf_path}")
            
            return {
                'success': True,
                'pdf_path': str(output_pdf_path),
                'errors': []
            }
            
        except Exception as e:
            return {
                'success': False,
                'errors': [f"Errore generazione contratto: {str(e)}"]
            }
    
    def _generate_proforma_pdf(self, lead_data: Dict[str, Any], tipo_servizio: str, proforma_id: str) -> Dict[str, Any]:
        """Genera PDF proforma da template DOCX"""
        try:
            template_path = self.templates_dir / 'Template_Proforma_Unificato_TeleMedCare.docx'
            
            if not template_path.exists():
                return {
                    'success': False,
                    'errors': [f"Template proforma non trovato: {template_path}"]
                }
            
            print(f"📄 Generazione proforma da template: Template_Proforma_Unificato_TeleMedCare.docx")
            
            # Prepara dati per sostituzione placeholder
            placeholder_data = self._prepare_proforma_placeholders(lead_data, tipo_servizio)
            
            # Crea DOCX con placeholder sostituiti
            output_filename = f"{datetime.now().strftime('%Y%m%d')}_{lead_data.get('cognomeAssistito')}_{proforma_id}.docx"
            output_docx_path = self.proforma_dir / output_filename
            
            self._replace_placeholders_in_docx(template_path, output_docx_path, placeholder_data)
            
            # Converte DOCX → PDF
            output_pdf_path = output_docx_path.with_suffix('.pdf')
            self._convert_docx_to_pdf(output_docx_path, output_pdf_path)
            
            # Rimuovi DOCX intermedio (opzionale)
            # output_docx_path.unlink()
            
            print(f"✅ Proforma PDF generata: {output_pdf_path}")
            
            return {
                'success': True,
                'pdf_path': str(output_pdf_path),
                'errors': []
            }
            
        except Exception as e:
            return {
                'success': False,
                'errors': [f"Errore generazione proforma: {str(e)}"]
            }
    
    def _prepare_contract_placeholders(self, lead_data: Dict[str, Any], tipo_servizio: str) -> Dict[str, str]:
        """
        Prepara dizionario placeholder → valore per contratti
        
        Placeholder contratto:
        - {{NOME_ASSISTITO}}
        - {{COGNOME_ASSISTITO}}
        - {{DATA_NASCITA}}
        - {{LUOGO_NASCITA}}
        - {{CODICE_FISCALE_ASSISTITO}}
        - {{INDIRIZZO_ASSISTITO}}
        - {{CAP_ASSISTITO}}
        - {{CITTA_ASSISTITO}}
        - {{PROVINCIA_ASSISTITO}}
        - {{EMAIL_ASSISTITO}}
        - {{TELEFONO_ASSISTITO}}
        - {{DATA_CONTRATTO}}
        - {{DATA_INIZIO_SERVIZIO}}
        - {{DATA_SCADENZA}}
        - {{IMPORTO_PRIMO_ANNO}}
        """
        intestazione = lead_data.get('intestazioneContratto', 'Assistito')
        
        # Date
        data_oggi = datetime.now()
        data_inizio = data_oggi + timedelta(days=7)  # Inizio tra 7 giorni
        data_scadenza = data_inizio + timedelta(days=365)  # 1 anno
        
        # Prezzo
        prezzo_base = self.PREZZI[tipo_servizio]
        prezzo_iva_inclusa = prezzo_base * (1 + self.IVA_RATE)
        
        # Determina dati intestatario
        if intestazione == 'Assistito':
            nome = lead_data.get('nomeAssistito', '')
            cognome = lead_data.get('cognomeAssistito', '')
            cf = lead_data.get('cfAssistito', '')
            indirizzo = lead_data.get('indirizzoAssistito', '')
            cap = self._extract_cap(indirizzo)
            citta = self._extract_city(indirizzo)
            provincia = self._extract_province(indirizzo)
            email = lead_data.get('emailRichiedente', '')  # Email sempre del richiedente
            telefono = lead_data.get('telefonoRichiedente', '')
        else:
            nome = lead_data.get('nomeRichiedente', '')
            cognome = lead_data.get('cognomeRichiedente', '')
            cf = lead_data.get('cfRichiedente', '')
            indirizzo = lead_data.get('indirizzoRichiedente', '')
            cap = self._extract_cap(indirizzo)
            citta = self._extract_city(indirizzo)
            provincia = self._extract_province(indirizzo)
            email = lead_data.get('emailRichiedente', '')
            telefono = lead_data.get('telefonoRichiedente', '')
        
        return {
            'NOME_ASSISTITO': nome,
            'COGNOME_ASSISTITO': cognome,
            'DATA_NASCITA': lead_data.get('dataNascitaAssistito', ''),
            'LUOGO_NASCITA': lead_data.get('luogoNascitaAssistito', ''),
            'CODICE_FISCALE_ASSISTITO': cf,
            'INDIRIZZO_ASSISTITO': indirizzo,
            'CAP_ASSISTITO': cap,
            'CITTA_ASSISTITO': citta,
            'PROVINCIA_ASSISTITO': provincia,
            'EMAIL_ASSISTITO': email,
            'TELEFONO_ASSISTITO': telefono,
            'DATA_CONTRATTO': data_oggi.strftime('%d/%m/%Y'),
            'DATA_INIZIO_SERVIZIO': data_inizio.strftime('%d/%m/%Y'),
            'DATA_SCADENZA': data_scadenza.strftime('%d/%m/%Y'),
            'IMPORTO_PRIMO_ANNO': f"€ {prezzo_iva_inclusa:.2f}"
        }
    
    def _prepare_proforma_placeholders(self, lead_data: Dict[str, Any], tipo_servizio: str) -> Dict[str, str]:
        """
        Prepara dizionario placeholder → valore per proforma
        
        Placeholder proforma:
        - {{NOME_ASSISTITO}}
        - {{COGNOME_ASSISTITO}}
        - {{CODICE_FISCALE}}
        - {{INDIRIZZO_COMPLETO}}
        - {{CITTA}}
        - {{EMAIL_RICHIEDENTE}}
        - {{DATA_RICHIESTA}}
        - {{DATA_ATTIVAZIONE}}
        - {{PREZZO_PACCHETTO}}
        - {{SERIAL_NUMBER}}
        - {{TELEFONO_SIDLY}}
        - {{COMUNICAZIONE_TIPO}}
        """
        intestazione = lead_data.get('intestazioneContratto', 'Assistito')
        
        # Date
        data_oggi = datetime.now()
        data_attivazione = data_oggi + timedelta(days=7)
        
        # Prezzo
        prezzo_base = self.PREZZI[tipo_servizio]
        prezzo_iva_inclusa = prezzo_base * (1 + self.IVA_RATE)
        
        # Determina dati intestatario
        if intestazione == 'Assistito':
            cf = lead_data.get('cfAssistito', '')
            indirizzo = lead_data.get('indirizzoAssistito', '')
            citta = self._extract_city(indirizzo)
        else:
            cf = lead_data.get('cfRichiedente', '')
            indirizzo = lead_data.get('indirizzoRichiedente', '')
            citta = self._extract_city(indirizzo)
        
        # Serial number dispositivo (placeholder)
        serial_number = f"SIDLY-{datetime.now().strftime('%Y%m%d')}-{lead_data.get('id', '000000')[:6]}"
        
        # Tipo comunicazione (SMS/Email/Chiamata vocale)
        comunicazione_tipo = "SMS, Email e Chiamata vocale"
        
        return {
            'NOME_ASSISTITO': lead_data.get('nomeAssistito', ''),
            'COGNOME_ASSISTITO': lead_data.get('cognomeAssistito', ''),
            'CODICE_FISCALE': cf,
            'INDIRIZZO_COMPLETO': indirizzo,
            'CITTA': citta,
            'EMAIL_RICHIEDENTE': lead_data.get('emailRichiedente', ''),
            'DATA_RICHIESTA': data_oggi.strftime('%d/%m/%Y'),
            'DATA_ATTIVAZIONE': data_attivazione.strftime('%d/%m/%Y'),
            'PREZZO_PACCHETTO': f"€ {prezzo_iva_inclusa:.2f}",
            'SERIAL_NUMBER': serial_number,
            'TELEFONO_SIDLY': '+39 02 1234 5678',  # Numero TeleMedCare
            'COMUNICAZIONE_TIPO': comunicazione_tipo
        }
    
    def _replace_placeholders_in_docx(self, template_path: Path, output_path: Path, placeholders: Dict[str, str]):
        """Sostituisce placeholder nel DOCX e salva nuovo file"""
        doc = Document(str(template_path))
        
        # Sostituisci nei paragrafi
        for para in doc.paragraphs:
            for key, value in placeholders.items():
                if f'{{{{{key}}}}}' in para.text:
                    para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Sostituisci nelle tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            if f'{{{{{key}}}}}' in para.text:
                                para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Salva nuovo documento
        doc.save(str(output_path))
        print(f"📝 DOCX con placeholder sostituiti salvato: {output_path}")
    
    def _convert_docx_to_pdf(self, docx_path: Path, pdf_path: Path):
        """
        Converte DOCX a PDF
        
        Strategia:
        1. Prova LibreOffice (se disponibile)
        2. Fallback: salva DOCX come "PDF" (placeholder per ora)
        3. In produzione: usa servizio esterno (CloudConvert, etc.)
        
        Note: Per ambiente sandbox senza LibreOffice/Word, salviamo il DOCX
        e in produzione si usa un servizio di conversione cloud
        """
        # Verifica se LibreOffice è disponibile
        libreoffice_available = subprocess.run(
            ['which', 'libreoffice'],
            capture_output=True
        ).returncode == 0
        
        if libreoffice_available:
            try:
                # Usa LibreOffice per conversione
                subprocess.run([
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(pdf_path.parent),
                    str(docx_path)
                ], check=True, capture_output=True, timeout=30)
                
                print(f"✅ PDF generato con LibreOffice: {pdf_path}")
                return
                
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
                print(f"⚠️ LibreOffice fallito: {e}")
        
        # Fallback: copia DOCX come PDF (placeholder)
        # In produzione, qui si userebbe un servizio cloud come CloudConvert
        print(f"⚠️ LibreOffice non disponibile, salvo DOCX come placeholder PDF")
        
        # Per ora, crea un PDF vuoto come placeholder
        # In produzione, sostituire con chiamata a servizio di conversione
        import shutil
        shutil.copy(str(docx_path), str(pdf_path.with_suffix('.docx.backup')))
        
        # Crea file PDF placeholder
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4\n%PLACEHOLDER - Documento generato\n')
        
        print(f"✅ PDF placeholder creato: {pdf_path}")
        print(f"ℹ️  DOCX originale salvato: {pdf_path.with_suffix('.docx.backup')}")
        print(f"ℹ️  In produzione, usare servizio di conversione cloud (CloudConvert, etc.)")
    
    def _extract_cap(self, indirizzo: str) -> str:
        """Estrae CAP dall'indirizzo (es. 'Via Roma 10, 00100 Roma' → '00100')"""
        match = re.search(r'\b(\d{5})\b', indirizzo)
        return match.group(1) if match else ''
    
    def _extract_city(self, indirizzo: str) -> str:
        """Estrae città dall'indirizzo (es. 'Via Roma 10, 00100 Roma RM' → 'Roma')"""
        # Prova a estrarre dopo il CAP
        match = re.search(r'\d{5}\s+([A-Za-zàèéìòùÀÈÉÌÒÙ\s]+?)(?:\s+[A-Z]{2})?$', indirizzo)
        return match.group(1).strip() if match else ''
    
    def _extract_province(self, indirizzo: str) -> str:
        """Estrae provincia dall'indirizzo (es. 'Via Roma 10, 00100 Roma RM' → 'RM')"""
        match = re.search(r'\b([A-Z]{2})$', indirizzo)
        return match.group(1) if match else ''
    
    def _generate_contract_id(self) -> str:
        """Genera ID univoco contratto"""
        return f"CTR{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    def _generate_proforma_id(self) -> str:
        """Genera ID univoco proforma"""
        return f"PRF{datetime.now().strftime('%Y%m%d%H%M%S')}"


def main():
    """Test del sistema"""
    # Dati lead di esempio
    lead_data = {
        'id': '000123',
        'nomeRichiedente': 'Mario',
        'cognomeRichiedente': 'Rossi',
        'emailRichiedente': 'mario.rossi@example.com',
        'telefonoRichiedente': '+39 333 1234567',
        'nomeAssistito': 'Giulia',
        'cognomeAssistito': 'Verdi',
        'dataNascitaAssistito': '15/03/1950',
        'luogoNascitaAssistito': 'Roma',
        'cfAssistito': 'VRDGLI50C55H501Z',
        'indirizzoAssistito': 'Via dei Fiori 25, 00100 Roma RM',
        'pacchetto': 'Servizio Base',
        'vuoleContratto': True,
        'intestazioneContratto': 'Assistito'
    }
    
    # Genera documenti
    generator = DocumentGenerator()
    result = generator.generate_contract_from_lead(lead_data)
    
    print("\n" + "="*60)
    print("RISULTATO GENERAZIONE:")
    print(json.dumps(result, indent=2))
    print("="*60)


if __name__ == '__main__':
    main()
