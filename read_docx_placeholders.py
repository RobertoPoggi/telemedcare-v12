#!/usr/bin/env python3
"""
Script per estrarre placeholder dai template DOCX
"""
import sys
from docx import Document
import re

def extract_placeholders(docx_path):
    """Estrae tutti i placeholder dal documento DOCX"""
    doc = Document(docx_path)
    placeholders = set()
    
    # Cerca nei paragrafi
    for para in doc.paragraphs:
        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
        placeholders.update(matches)
    
    # Cerca nelle tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                    placeholders.update(matches)
    
    return sorted(placeholders)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python read_docx_placeholders.py <file.docx>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    placeholders = extract_placeholders(docx_path)
    
    print(f"\n📄 Template: {docx_path}")
    print(f"🔍 Trovati {len(placeholders)} placeholder:\n")
    for placeholder in placeholders:
        print(f"  - {{{{{placeholder}}}}}")
