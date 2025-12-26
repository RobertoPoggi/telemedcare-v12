#!/usr/bin/env python3
"""
DOCX Contract Generator - TeleMedCare V11.0
Genera contratti PDF da template DOCX sostituendo i placeholder
"""

from docx import Document
import sys
import json
import os
from datetime import datetime

def replace_placeholder_in_paragraph(paragraph, placeholders):
    """Sostituisce placeholder in un paragrafo mantenendo la formattazione"""
    full_text = paragraph.text
    
    for key, value in placeholders.items():
        placeholder = f"{{{{{key}}}}}"
        
        if placeholder in full_text:
            # Il placeholder è presente nel testo completo
            # Ricostruisci il paragrafo sostituendo il placeholder
            
            # Salva lo stile del primo run
            if len(paragraph.runs) > 0:
                first_run_font = paragraph.runs[0].font
            
            # Sostituisci nel testo completo
            new_text = full_text.replace(placeholder, value)
            
            # Pulisci tutti i runs esistenti
            for run in paragraph.runs:
                run.text = ''
            
            # Crea un nuovo run con il testo sostituito
            new_run = paragraph.runs[0] if len(paragraph.runs) > 0 else paragraph.add_run()
            new_run.text = new_text
            
            # Aggiorna il full_text per prossime sostituzioni
            full_text = new_text

def replace_placeholders_in_docx(template_path, placeholders, output_path):
    """
    Sostituisce tutti i placeholder nel template DOCX
    
    Args:
        template_path: Percorso template DOCX
        placeholders: Dict con placeholder → valore
        output_path: Percorso output DOCX
    """
    print(f"📄 Caricamento template: {template_path}")
    
    # Carica il documento
    doc = Document(template_path)
    
    # Sostituisci placeholder nei paragrafi
    for para in doc.paragraphs:
        replace_placeholder_in_paragraph(para, placeholders)
    
    # Sostituisci placeholder nelle tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholder_in_paragraph(para, placeholders)
    
    # Salva il documento modificato
    print(f"💾 Salvataggio contratto: {output_path}")
    doc.save(output_path)
    print(f"✅ Contratto DOCX generato con successo")
    
    return output_path

def docx_to_pdf_cloudflare(docx_path, output_dir='/tmp'):
    """
    Conversione DOCX → PDF usando Cloudflare Browser Rendering
    (alternativa a LibreOffice)
    
    Nota: Questa funzione deve essere chiamata dal Worker con env.BROWSER
    """
    print("ℹ️  Conversione DOCX → PDF richiede Cloudflare Browser Rendering")
    print("ℹ️  In sviluppo locale, usa il DOCX direttamente")
    
    # Per ora, ritorna None
    # La conversione PDF sarà gestita dal Browser Rendering in produzione
    return None

if __name__ == '__main__':
    # Leggi i dati da stdin (JSON)
    if len(sys.argv) > 1:
        data = json.loads(sys.argv[1])
    else:
        data = json.load(sys.stdin)
    
    template_path = data.get('template_path', '/home/user/webapp/templates/contracts/Template_Contratto_eCura.docx')
    output_path = data.get('output_path', '/tmp/contratto_generato.docx')
    placeholders = data.get('placeholders', {})
    
    # Genera il contratto
    result = replace_placeholders_in_docx(template_path, placeholders, output_path)
    
    # Output JSON con il percorso del file generato
    output = {
        'success': True,
        'docx_path': result,
        'pdf_path': None  # PDF sarà generato in produzione
    }
    
    print(json.dumps(output))
